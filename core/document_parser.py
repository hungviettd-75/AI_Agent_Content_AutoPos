"""core/document_parser.py
===========================
Module trích xuất văn bản thuần từ nhiều định dạng tài liệu:
  - PDF       (.pdf)   → pdfplumber
  - DOCX      (.docx)  → python-docx
  - TXT       (.txt)   → đọc thẳng
  - Markdown  (.md)    → đọc thẳng
  - Website   (URL)    → requests + BeautifulSoup

Tất cả hàm trả về dict:
  {
    "text"    : str,          # Văn bản thô đầy đủ
    "preview" : str,          # 500 ký tự đầu
    "pages"   : int,          # Số trang (nếu PDF)
    "word_count": int,
    "error"   : str | None,   # None nếu thành công
  }
"""

import io
import re
from pathlib import Path
from config.config import logger


# ─── Helpers ────────────────────────────────────────────────────────────────
def _result(text: str, pages: int = 1, error: str | None = None) -> dict:
    clean = re.sub(r"\s{3,}", "\n\n", text.strip())
    return {
        "text":       clean,
        "preview":    clean[:600] + ("..." if len(clean) > 600 else ""),
        "pages":      pages,
        "word_count": len(clean.split()),
        "error":      error,
    }


# ─── PDF ────────────────────────────────────────────────────────────────────
def parse_pdf(file_bytes: bytes) -> dict:
    """Trích xuất văn bản từ file PDF (bytes)."""
    try:
        import pdfplumber
        texts = []
        pages = 0
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        return _result("\n\n".join(texts), pages=pages)
    except ImportError:
        return _result("", error="Thiếu thư viện pdfplumber. Chạy: pip install pdfplumber")
    except Exception as e:
        logger.error(f"parse_pdf lỗi: {e}")
        return _result("", error=f"Không đọc được PDF: {e}")


# ─── DOCX ───────────────────────────────────────────────────────────────────
def parse_docx(file_bytes: bytes) -> dict:
    """Trích xuất văn bản từ file DOCX (bytes)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Thêm text từ tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return _result("\n\n".join(paragraphs))
    except ImportError:
        return _result("", error="Thiếu thư viện python-docx. Chạy: pip install python-docx")
    except Exception as e:
        logger.error(f"parse_docx lỗi: {e}")
        return _result("", error=f"Không đọc được DOCX: {e}")


# ─── TXT ────────────────────────────────────────────────────────────────────
def parse_txt(file_bytes: bytes) -> dict:
    """Đọc văn bản từ file TXT (thử UTF-8, fallback UTF-16, latin-1)."""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(enc)
            return _result(text)
        except (UnicodeDecodeError, ValueError):
            continue
    return _result("", error="Không thể giải mã file TXT (encoding không hỗ trợ)")


# ─── Markdown ───────────────────────────────────────────────────────────────
def parse_markdown(file_bytes: bytes) -> dict:
    """Đọc Markdown, giữ nguyên nội dung gốc."""
    try:
        text = file_bytes.decode("utf-8")
        # Loại bỏ cú pháp Markdown thô để đếm từ dễ hơn
        clean = re.sub(r"#{1,6}\s", "", text)    # headers
        clean = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", clean)  # bold/italic
        clean = re.sub(r"`(.+?)`", r"\1", clean)  # inline code
        clean = re.sub(r"!\[.*?\]\(.*?\)", "", clean)  # images
        clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)  # links
        return _result(clean)
    except Exception as e:
        logger.error(f"parse_markdown lỗi: {e}")
        return _result("", error=f"Không đọc được Markdown: {e}")


# ─── Website URL ────────────────────────────────────────────────────────────
def parse_url(url: str, timeout: int = 15) -> dict:
    """
    Tải và trích xuất nội dung văn bản từ URL (HTML → text thuần).
    Loại bỏ script, style, nav, footer để giữ nội dung chính.
    """
    try:
        import requests
        from bs4 import BeautifulSoup

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            )
        }
        resp = requests.get(url, headers=headers, timeout=timeout)
        resp.raise_for_status()
        resp.encoding = resp.apparent_encoding or "utf-8"

        soup = BeautifulSoup(resp.text, "html.parser")

        # Xóa các phần không cần thiết
        for tag in soup(["script", "style", "nav", "footer",
                          "header", "aside", "form", "button",
                          "noscript", "iframe", "svg", "img"]):
            tag.decompose()

        # Ưu tiên lấy từ thẻ article/main nếu có
        main_content = (
            soup.find("article")
            or soup.find("main")
            or soup.find(id=re.compile(r"content|main|body", re.I))
            or soup.find("body")
            or soup
        )

        # Lấy text từng đoạn, giữ nguyên dòng
        lines = []
        for el in main_content.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "blockquote"]
        ):
            t = el.get_text(" ", strip=True)
            if t and len(t) > 20:
                lines.append(t)

        if not lines:
            lines = [main_content.get_text(" ", strip=True)]

        text = "\n\n".join(lines)
        return _result(text if text.strip() else "Không lấy được nội dung từ URL này.")

    except ImportError:
        return _result("", error="Thiếu thư viện requests hoặc beautifulsoup4.")
    except Exception as e:
        logger.error(f"parse_url '{url}' lỗi: {e}")
        return _result("", error=f"Lỗi tải URL: {e}")


# ─── Dispatcher chính ────────────────────────────────────────────────────────
def parse_document(
    file_bytes: bytes | None = None,
    filename: str = "",
    url: str = "",
) -> dict:
    """
    Gọi parser phù hợp dựa trên đuôi file hoặc URL.

    Args:
        file_bytes : Nội dung file (bytes) — dùng cho PDF/DOCX/TXT/MD
        filename   : Tên file (để xác định định dạng)
        url        : URL website — nếu có, ưu tiên parse URL

    Returns:
        dict với keys: text, preview, pages, word_count, error
    """
    if url:
        return parse_url(url)

    if not file_bytes or not filename:
        return _result("", error="Không có dữ liệu để xử lý.")

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    elif ext in (".txt",):
        return parse_txt(file_bytes)
    elif ext in (".md", ".markdown"):
        return parse_markdown(file_bytes)
    else:
        return _result("", error=f"Định dạng '{ext}' chưa được hỗ trợ.")
