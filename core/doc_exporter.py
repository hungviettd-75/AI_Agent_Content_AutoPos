import os
import io
import requests
import unicodedata
import pandas as pd
from datetime import datetime
from fpdf import FPDF

def _knowledge_value(row, key, default=""):
    try:
        value = row.get(key, default)
    except AttributeError:
        try:
            value = row[key]
        except Exception:
            value = default
    if value is None:
        return default
    return value

def get_plan_items(plan_data):
    if isinstance(plan_data, dict):
        items = plan_data.get("items", [])
        return items if isinstance(items, list) else []
    return plan_data if isinstance(plan_data, list) else []

def export_post_to_pdf(topic, platform, content, image_prompts=""):
    try:
        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.add_page()
        
        font_paths = ["DejaVuSans.ttf", "C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\tahoma.ttf"]
        if not os.path.exists("DejaVuSans.ttf"):
            try:
                url = "https://raw.githubusercontent.com/halfmoon-tech/dejavu-fonts-ttf/master/DejaVuSans.ttf"
                r = requests.get(url, timeout=5)
                if r.status_code == 200:
                    with open("DejaVuSans.ttf", "wb") as f: f.write(r.content)
            except: pass
        
        font_loaded = False
        for fpath in font_paths:
            if os.path.exists(fpath):
                try:
                    pdf.add_font("CustomFont", "", fpath)
                    pdf.set_font("CustomFont", size=12)
                    font_loaded = True
                    break
                except: continue
        if not font_loaded: pdf.set_font("Arial", size=12)

        # Header
        if font_loaded:
            pdf.set_font("CustomFont", size=16)
        else:
            pdf.set_font("Arial", style="B", size=16)
        pdf.cell(0, 10, txt="BÀI VIẾT MARKETING AI", ln=True, align='C')
        pdf.ln(10)

        # Details
        pdf.set_font("CustomFont" if font_loaded else "Arial", size=11)
        pdf.set_x(10)
        pdf.multi_cell(190, 8, txt=f"Chủ đề: {topic}")
        pdf.set_x(10)
        pdf.multi_cell(190, 8, txt=f"Nền tảng: {platform}")
        pdf.set_x(10)
        pdf.multi_cell(190, 8, txt=f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}")
        pdf.ln(5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(10)

        # Content
        pdf.set_font("CustomFont" if font_loaded else "Arial", size=12)
        lines = content.split('\n')
        for line in lines:
            try:
                pdf.set_x(10)
                pdf.multi_cell(190, 7, txt=line)
            except:
                clean_line = "".join(c for c in unicodedata.normalize('NFKD', line) if not unicodedata.combining(c))
                pdf.set_x(10)
                pdf.multi_cell(190, 7, txt=clean_line)
        
        if image_prompts:
            pdf.ln(10)
            if font_loaded:
                pdf.set_font("CustomFont", size=12)
            else:
                pdf.set_font("Arial", style="B", size=12)
            pdf.cell(0, 10, txt="Gợi ý Prompt vẽ ảnh AI:", ln=True)
            pdf.set_font("CustomFont" if font_loaded else "Arial", size=10)
            pdf.set_x(10)
            pdf.multi_cell(190, 6, txt=image_prompts)

        return bytes(pdf.output())
    except Exception as e:
        import traceback
        print(f"Lỗi tạo nội dung PDF: {e}")
        print(traceback.format_exc())
        return None

def export_all_posts_to_pdf(dataframe):
    try:
        pdf = FPDF(orientation="L", unit="mm", format="A4")
        pdf.add_page()
        
        font_paths = ["DejaVuSans.ttf", "C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\tahoma.ttf"]
        if not os.path.exists("DejaVuSans.ttf"):
            try:
                url = "https://raw.githubusercontent.com/halfmoon-tech/dejavu-fonts-ttf/master/DejaVuSans.ttf"
                r = requests.get(url, timeout=5)
                if r.status_code == 200:
                    with open("DejaVuSans.ttf", "wb") as f: f.write(r.content)
            except: pass
        
        font_loaded = False
        for fpath in font_paths:
            if os.path.exists(fpath):
                try:
                    pdf.add_font("CustomFont", "", fpath)
                    pdf.set_font("CustomFont", size=10)
                    font_loaded = True
                    break
                except: continue
        if not font_loaded: pdf.set_font("Arial", size=10)

        # Title
        if font_loaded:
            pdf.set_font("CustomFont", size=14)
        else:
            pdf.set_font("Arial", style="B", size=14)
        pdf.cell(0, 10, txt="BÁO CÁO LỊCH SỬ BÀI VIẾT MARKETING AI", ln=True, align='C')
        pdf.ln(5)

        # Table Header
        col_widths = [15, 40, 30, 140, 50]
        cols = ['ID', 'DATE', 'PLATFORM', 'TOPIC', 'STATUS']
        pdf.set_fill_color(230, 230, 230)
        if font_loaded:
            pdf.set_font("CustomFont", size=10)
        else:
            pdf.set_font("Arial", style="B", size=10)
        for i, col in enumerate(cols):
            pdf.cell(col_widths[i], 10, txt=col, border=1, fill=True, align='C')
        pdf.ln()

        # Table Rows
        pdf.set_font("CustomFont" if font_loaded else "Arial", size=9)
        cols_to_print = ['id', 'date', 'platform', 'topic', 'status']
        for _, row in dataframe.iterrows():
            for i, col in enumerate(cols_to_print):
                val = str(row[col]).replace('\n', ' ')
                if col == 'topic' and len(val) > 70: val = val[:67] + "..."
                elif col != 'topic' and len(val) > 25: val = val[:22] + "..."
                
                try:
                    pdf.cell(col_widths[i], 10, txt=val, border=1)
                except:
                    clean_val = "".join(c for c in unicodedata.normalize('NFKD', val) if not unicodedata.combining(c))
                    pdf.cell(col_widths[i], 10, txt=clean_val, border=1)
            pdf.ln()

        return bytes(pdf.output())
    except Exception as e:
        import traceback
        print(f"Lỗi tạo báo cáo PDF: {e}")
        print(traceback.format_exc())
        return None

def export_knowledge_to_pdf(row):
    try:
        def clean(value):
            return str(value or "")

        def safe_text(value):
            value = clean(value)
            if font_loaded:
                return value
            return "".join(ch for ch in unicodedata.normalize("NFKD", value) if not unicodedata.combining(ch))

        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=18)
        pdf.add_page()

        font_paths = ["C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\tahoma.ttf", "DejaVuSans.ttf"]
        font_loaded = False
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdf.add_font("KnowledgeFont", "", font_path)
                    pdf.add_font("KnowledgeFont", "B", font_path)
                    font_loaded = True
                    break
                except Exception:
                    continue

        font_name = "KnowledgeFont" if font_loaded else "Arial"

        # Watermark
        pdf.set_text_color(232, 238, 247)
        pdf.set_font(font_name, "B", 28)
        pdf.text(58, 150, safe_text("hungvietai.com"))

        # Logo block
        pdf.set_fill_color(30, 58, 138)
        pdf.rect(12, 12, 42, 14, style="F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(font_name, "B", 10)
        pdf.set_xy(12, 16)
        pdf.cell(42, 6, safe_text("HUNG VIET AI"), align="C")

        pdf.set_text_color(30, 58, 138)
        pdf.set_font(font_name, "B", 18)
        pdf.set_xy(12, 34)
        pdf.multi_cell(186, 9, safe_text("AI Knowledge Sharing"), align="C")

        topic = clean(_knowledge_value(row, "topic", "AI Knowledge"))
        metadata = [
            ("Topic", topic),
            ("Date", _knowledge_value(row, "date")),
            ("Platform", _knowledge_value(row, "platform")),
            ("AI Tool", _knowledge_value(row, "tool_name")),
            ("Audience", _knowledge_value(row, "audience")),
            ("Difficulty", _knowledge_value(row, "difficulty")),
            ("Knowledge Type", _knowledge_value(row, "knowledge_type")),
            ("Status", _knowledge_value(row, "status")),
        ]

        pdf.ln(8)
        pdf.set_fill_color(239, 246, 255)
        pdf.set_draw_color(191, 219, 254)
        pdf.set_text_color(15, 23, 42)
        pdf.set_font(font_name, "", 10)
        for label, value in metadata:
            pdf.set_x(18)
            pdf.set_font(font_name, "B", 10)
            pdf.cell(38, 8, safe_text(label), border=1, fill=True)
            pdf.set_font(font_name, "", 10)
            pdf.multi_cell(136, 8, safe_text(value), border=1)

        pdf.ln(6)
        pdf.set_text_color(30, 58, 138)
        pdf.set_font(font_name, "B", 13)
        pdf.cell(0, 8, safe_text("Content"), ln=True)
        pdf.set_draw_color(37, 99, 235)
        pdf.line(18, pdf.get_y(), 192, pdf.get_y())
        pdf.ln(4)

        pdf.set_text_color(30, 41, 59)
        for line in clean(_knowledge_value(row, "content")).split("\n"):
            stripped = line.strip()
            if stripped.startswith("###") or stripped.startswith("##"):
                pdf.ln(2)
                pdf.set_text_color(30, 58, 138)
                pdf.set_font(font_name, "B", 12)
                pdf.multi_cell(174, 7, safe_text(stripped.replace("#", "").strip()))
                pdf.set_text_color(30, 41, 59)
                pdf.set_font(font_name, "", 10)
            elif stripped:
                pdf.set_font(font_name, "", 10)
                pdf.multi_cell(174, 6, safe_text(stripped))
            else:
                pdf.ln(2)

        pdf.set_y(-16)
        pdf.set_text_color(100, 116, 139)
        pdf.set_font(font_name, "", 8)
        pdf.cell(0, 8, safe_text("hungvietai.com | AI Knowledge Export"), align="C")
        return bytes(pdf.output())
    except Exception as e:
        import traceback
        print(f"Lỗi tạo PDF AI Knowledge: {e}")
        print(traceback.format_exc())
        return None

def export_knowledge_to_docx(row):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def shade_cell(cell, fill):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        title = doc.add_heading("AI Knowledge Sharing", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)
            run.font.size = Pt(20)

        subtitle = doc.add_paragraph(str(_knowledge_value(row, "topic", "AI Knowledge")))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.color.rgb = RGBColor(71, 85, 105)
            run.font.size = Pt(11)

        metadata = [
            ("Date", _knowledge_value(row, "date")),
            ("Platform", _knowledge_value(row, "platform")),
            ("AI Tool", _knowledge_value(row, "tool_name")),
            ("Audience", _knowledge_value(row, "audience")),
            ("Difficulty", _knowledge_value(row, "difficulty")),
            ("Knowledge Type", _knowledge_value(row, "knowledge_type")),
            ("Status", _knowledge_value(row, "status")),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "Metadata"
        table.rows[0].cells[1].text = "Value"
        for cell in table.rows[0].cells:
            shade_cell(cell, "1E3A8A")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for label, value in metadata:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value or "")
            shade_cell(cells[0], "EFF6FF")

        doc.add_paragraph("")
        doc.add_heading("Content", level=1)
        for raw_line in str(_knowledge_value(row, "content")).split("\n"):
            line = raw_line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("###") or line.startswith("##"):
                doc.add_heading(line.replace("#", "").strip(), level=2)
            else:
                paragraph = doc.add_paragraph(line)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(30, 41, 59)

        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = "hungvietai.com | AI Knowledge Export"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 116, 139)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        import traceback
        print(f"Lỗi tạo Word AI Knowledge: {e}")
        print(traceback.format_exc())
        return None

def export_knowledge_to_csv(dataframe):
    metadata_columns = [
        "id", "date", "platform", "topic", "audience", "tool_name",
        "knowledge_type", "difficulty", "summary", "status", "content"
    ]
    export_df = dataframe.copy()
    for column in metadata_columns:
        if column not in export_df.columns:
            export_df[column] = ""
    return export_df[metadata_columns].to_csv(index=False).encode("utf-8-sig")

def export_plan_to_word(plan_data):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def shade_cell(cell, fill):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        def set_cell_text_style(cell, bold=False, color=None, size=10):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = bold
                    run.font.size = Pt(size)
                    if color:
                        run.font.color.rgb = color

        plan_items = get_plan_items(plan_data)
        meta = plan_data.get("meta", {}) if isinstance(plan_data, dict) else {}

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        # Tiêu đề
        title = doc.add_heading('', level=0)
        run = title.add_run('📅 KẾ HOẠCH NỘI DUNG')
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(30, 58, 138)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thông tin ngày tạo file
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Ngày xuất file: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(100, 116, 139)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata kế hoạch
        if meta:
            doc.add_paragraph('')
            meta_heading = doc.add_heading('Thông tin kế hoạch', level=1)
            for run in meta_heading.runs:
                run.font.color.rgb = RGBColor(30, 58, 138)

            meta_labels = [
                ('topic', 'Chủ đề chính'),
                ('target', 'Đối tượng mục tiêu'),
                ('goal', 'Mục tiêu nội dung'),
                ('days', 'Số ngày'),
                ('style', 'Phong cách nội dung'),
                ('created_at', 'Ngày tạo'),
            ]
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = 'Light Grid Accent 1'
            meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            meta_table.rows[0].cells[0].text = 'Trường'
            meta_table.rows[0].cells[1].text = 'Thông tin'
            for cell in meta_table.rows[0].cells:
                shade_cell(cell, '1E3A8A')
                set_cell_text_style(cell, bold=True, color=RGBColor(255, 255, 255), size=10)

            for key, label in meta_labels:
                if key in meta and meta.get(key) not in [None, '']:
                    cells = meta_table.add_row().cells
                    cells[0].text = label
                    cells[1].text = str(meta.get(key, ''))
                    shade_cell(cells[0], 'EFF6FF')
                    set_cell_text_style(cells[0], bold=True, color=RGBColor(30, 58, 138), size=10)
                    set_cell_text_style(cells[1], size=10)

        doc.add_paragraph('')
        plan_heading = doc.add_heading('Bảng kế hoạch', level=1)
        for run in plan_heading.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)

        supported_columns = [
            'day', 'topic', 'target', 'goal', 'format', 'angle', 'hook', 'cta'
        ]
        header_labels = {
            'day': 'Ngày',
            'topic': 'Chủ đề',
            'target': 'Đối tượng',
            'goal': 'Mục tiêu',
            'format': 'Định dạng',
            'angle': 'Góc triển khai',
            'hook': 'Hook',
            'cta': 'CTA',
        }
        headers = [key for key in supported_columns if any(key in item for item in plan_items)]
        if not headers:
            headers = list(plan_items[0].keys()) if plan_items else ['day', 'topic', 'target', 'angle']

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(headers):
            hdr_cells[i].text = header_labels.get(key, key.upper())
            shade_cell(hdr_cells[i], '1E3A8A')
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for item in plan_items:
            row_cells = table.add_row().cells
            for i, key in enumerate(headers):
                row_cells[i].text = str(item.get(key, ''))
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(30, 41, 59)

        # Footer
        doc.add_paragraph('')
        footer = doc.add_paragraph()
        footer_run = footer.add_run('🌐 Website: hungvietai.com | Tạo bởi AI-Agent Content & Auto-Post')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        import traceback
        print(f"Lỗi tạo file Word: {e}")
        print(traceback.format_exc())
        return None
